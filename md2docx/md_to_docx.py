#!/usr/bin/env python3
"""Convert a Markdown file to .docx with Korean-friendly font settings.

Usage:
  python md_to_word.py input.md [output.docx]
"""

from __future__ import annotations

import argparse
from pathlib import Path
from typing import Iterable, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn

try:
    import markdown
    from bs4 import BeautifulSoup, NavigableString
    from bs4.element import Tag
except ImportError as e:
    raise SystemExit(
        "필수 패키지가 없습니다: markdown, beautifulsoup4. "
        "pip install markdown beautifulsoup4 python-docx 로 설치 후 다시 실행하세요."
    ) from e


HEADING_STYLES = [
    "Heading 1",
    "Heading 2",
    "Heading 3",
    "Heading 4",
    "Heading 5",
    "Heading 6",
]


def set_eastasian_font(run_or_style, font_name: str) -> None:
    if hasattr(run_or_style, "_element") and run_or_style._element is not None:
        if hasattr(run_or_style._element, "rPr"):
            rpr = run_or_style._element.rPr
            if rpr is None:
                rpr = run_or_style._element.get_or_add_rPr()
            rpr.rFonts.set(qn("w:eastAsia"), font_name)


def apply_style_font(style, font_name: str, size: int = 11) -> None:
    style.font.name = font_name
    style.font.size = Pt(size)
    set_eastasian_font(style, font_name)


def configure_korean_fonts(doc: Document, korean_font: str, code_font: str) -> None:
    for style_name in ["Normal", "Body Text", "Quote", "Intense Quote"]:
        if style_name in doc.styles:
            apply_style_font(doc.styles[style_name], korean_font, 11)

    for style_name in HEADING_STYLES:
        if style_name in doc.styles:
            apply_style_font(doc.styles[style_name], korean_font, 16)

    if "List Bullet" in doc.styles:
        apply_style_font(doc.styles["List Bullet"], korean_font, 11)
    if "List Number" in doc.styles:
        apply_style_font(doc.styles["List Number"], korean_font, 11)

    for style_name in ["Code", "No Spacing"]:
        if style_name in doc.styles:
            apply_style_font(doc.styles[style_name], code_font, 10)


def add_inline(node, paragraph, default_font: str, code_font: str, bold: bool = False,
               italic: bool = False, underline: bool = False) -> None:
    for child in node.contents if isinstance(node, Tag) else [node]:
        if isinstance(child, NavigableString):
            text = str(child)
            if text:
                run = paragraph.add_run(text)
                run.bold = True if bold else None
                run.italic = True if italic else None
                run.underline = True if underline else None
                run.font.name = default_font
                run.font.size = Pt(11)
                set_eastasian_font(run, default_font)
            continue

        if not isinstance(child, Tag):
            continue

        tag = child.name.lower()
        if tag in {"strong", "b"}:
            add_inline(child, paragraph, default_font, code_font, bold=True, italic=italic, underline=underline)
        elif tag in {"em", "i"}:
            add_inline(child, paragraph, default_font, code_font, bold=bold, italic=True, underline=underline)
        elif tag == "u":
            add_inline(child, paragraph, default_font, code_font, bold=bold, italic=italic, underline=True)
        elif tag == "code":
            for code_child in child.contents:
                if isinstance(code_child, NavigableString):
                    run = paragraph.add_run(str(code_child))
                    run.font.name = code_font
                    run.font.size = Pt(10)
                    set_eastasian_font(run, code_font)
                    run.bold = True
        elif tag == "br":
            paragraph.add_run("\n")
        else:
            add_inline(child, paragraph, default_font, code_font, bold=bold, italic=italic, underline=underline)


def add_paragraph_from_node(document: Document, node, korean_font: str, code_font: str) -> None:
    if isinstance(node, NavigableString):
        text = str(node).strip()
        if text:
            paragraph = document.add_paragraph()
            add_inline(node, paragraph, korean_font, code_font)
        return

    if not isinstance(node, Tag):
        return

    tag = node.name.lower()

    if tag == "h1":
        document.add_heading(node.get_text("", strip=True), level=1)
    elif tag == "h2":
        document.add_heading(node.get_text("", strip=True), level=2)
    elif tag == "h3":
        document.add_heading(node.get_text("", strip=True), level=3)
    elif tag == "h4":
        document.add_heading(node.get_text("", strip=True), level=4)
    elif tag == "h5":
        document.add_heading(node.get_text("", strip=True), level=5)
    elif tag == "h6":
        document.add_heading(node.get_text("", strip=True), level=6)
    elif tag == "blockquote":
        paragraph = document.add_paragraph(style="Quote")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_inline(node, paragraph, korean_font, code_font)
    elif tag == "pre":
        text = node.get_text()
        paragraph = document.add_paragraph(style="No Spacing")
        run = paragraph.add_run(text)
        run.font.name = code_font
        run.font.size = Pt(10)
        set_eastasian_font(run, code_font)
    elif tag in {"ul", "ol"}:
        add_list_nodes(document, node, korean_font, code_font, ordered=(tag == "ol"))
    elif tag == "p":
        paragraph = document.add_paragraph()
        add_inline(node, paragraph, korean_font, code_font)
    elif tag == "hr":
        document.add_paragraph("-" * 16)
    elif tag == "table":
        # Very lightweight table support: keeps plain text per row.
        rows = node.find_all("tr")
        for row in rows:
            cells = [c.get_text(" ", strip=True) for c in row.find_all(["th", "td"])]
            if not cells:
                continue
            paragraph = document.add_paragraph()
            add_inline(NavigableString(" | ".join(cells)), paragraph, korean_font, code_font)
    else:
        # Inline-mixed or unknown block; best effort
        paragraph = document.add_paragraph()
        add_inline(node, paragraph, korean_font, code_font)


def add_list_nodes(document: Document, list_node: Tag, korean_font: str, code_font: str, ordered: bool, level: int = 0) -> None:
    li_nodes = list_node.find_all("li", recursive=False)
    for index, li in enumerate(li_nodes, start=1):
        paragraph = document.add_paragraph(style="List Number" if ordered else "List Bullet")
        if level:
            paragraph.paragraph_format.left_indent = Pt(level * 12)

        children: Iterable = list(li.children)
        for child in children:
            if isinstance(child, Tag) and child.name in {"ul", "ol"}:
                add_list_nodes(document, child, korean_font, code_font, ordered=(child.name == "ol"), level=level + 1)
            else:
                add_inline(child, paragraph, korean_font, code_font)


def convert_md_to_docx(input_path: Path, output_path: Path, korean_font: str, code_font: str) -> None:
    raw_md = input_path.read_text(encoding="utf-8")
    html = markdown.markdown(
        raw_md,
        extensions=[
            "extra",
            "sane_lists",
            "nl2br",
        ],
    )

    soup = BeautifulSoup(html, "html.parser")
    document = Document()
    configure_korean_fonts(document, korean_font, code_font)

    for node in soup.children:
        add_paragraph_from_node(document, node, korean_font, code_font)

    document.save(str(output_path))


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="md 파일을 docx로 변환")
    parser.add_argument("input", help="변환할 .md 파일")
    parser.add_argument("output", nargs="?", help="저장할 .docx 파일")
    parser.add_argument("--font", default="Malgun Gothic", help="본문 한글 기본 폰트")
    parser.add_argument("--code-font", default="Consolas", help="코드용 폰트")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    input_path = Path(args.input)
    if not input_path.exists():
        raise SystemExit(f"입력 파일이 없습니다: {input_path}")

    if input_path.suffix.lower() != ".md":
        raise SystemExit("입력 파일 확장자가 .md인지 확인하세요.")

    output_path = Path(args.output) if args.output else input_path.with_suffix(".docx")
    convert_md_to_docx(input_path, output_path, args.font, args.code_font)
    print(f"완료: {output_path}")


if __name__ == "__main__":
    main()
